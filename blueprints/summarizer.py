# blueprints/summarizer.py
import io
import os
import sys
import re
from flask import Blueprint, request, render_template, send_file, current_app, jsonify

# Sumy imports
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.lex_rank import LexRankSummarizer

# other parsers
import pdfplumber
import docx

# try to make sure nltk data from venv/share/nltk_data is visible
try:
    import nltk
    # typical venv share path relative to python executable
    venv_share = os.path.normpath(os.path.join(os.path.dirname(sys.executable), '..', 'share', 'nltk_data'))
    if os.path.isdir(venv_share) and venv_share not in nltk.data.path:
        nltk.data.path.append(venv_share)
except Exception:
    # silent fallback; we don't want this to break import if nltk not present here
    pass

summ_bp = Blueprint('summ_bp', __name__, url_prefix='/summarizer')

ALLOWED_TEXT_MIMES = {'text/plain'}
ALLOWED_FILE_EXT = {'txt', 'pdf', 'docx'}


def extract_text_from_pdf(file_stream):
    text_chunks = []
    # pdfplumber accepts a file-like object
    with pdfplumber.open(file_stream) as pdf:
        for p in pdf.pages:
            txt = p.extract_text()
            if txt:
                text_chunks.append(txt)
    return "\n".join(text_chunks)

def extract_text_from_docx(file_stream):
    try:
        doc = docx.Document(file_stream)
    except Exception:
        file_stream.seek(0)
        doc = docx.Document(io.BytesIO(file_stream.read()))

    pieces = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            pieces.append(text)

    if doc.tables:
        pieces.append("")
        for t_idx, table in enumerate(doc.tables):
            rows_text = []
            for r in table.rows:
                cells = []
                for c in r.cells:
                    cell_text_chunks = []
                    for p in c.paragraphs:
                        t = p.text.strip()
                        if t:
                            cell_text_chunks.append(t)
                    cell_text = " ".join(cell_text_chunks).strip()
                    cells.append(cell_text)
                rows_text.append(" | ".join([c for c in cells if c]))
            if rows_text:
                pieces.append("\n".join(rows_text))
            pieces.append("")

    combined = "\n\n".join([p for p in pieces if p is not None and str(p).strip() != ""])
    return combined

def simple_sentence_split(text):
    """
    Fallback simple sentence splitter.
    Splits on ., ?, ! followed by whitespace or end-of-string.
    Returns list of sentences (trimmed).
    """
    if not text or not text.strip():
        return []
    s = re.sub(r'\s+', ' ', text).strip()
    parts = re.split(r'(?<=[\.\?\!])\s+', s)
    parts = [p.strip() for p in parts if p and len(p) > 3]
    return parts


def summarize_text(text, sentence_count=5):
    """
    Summarize text using Sumy LexRank if possible.
    Fallbacks:
      - try Tokenizer("indonesian")
      - if missing, try Tokenizer("english")
      - if still missing or errors, use simple sentence split + heuristic ranking
    """
    if not text or not text.strip():
        return ""

    use_sumy = False
    parser = None

    # Attempt 1: Indonesian tokenizer (may not exist)
    try:
        parser = PlaintextParser.from_string(text, Tokenizer("indonesian"))
        use_sumy = True
    except Exception:
        current_app.logger.debug("Tokenizer('indonesian') failed, trying english tokenizer...")

    # Attempt 2: English tokenizer as fallback for sentence splitting
    if not use_sumy:
        try:
            parser = PlaintextParser.from_string(text, Tokenizer("english"))
            use_sumy = True
            current_app.logger.debug("Tokenizer('english') succeeded. Using english tokenizer as fallback.")
        except Exception:
            current_app.logger.debug("Tokenizer('english') also failed. Using simple fallback splitter.")

    # If we can use sumy, run LexRank
    if use_sumy and parser is not None:
        try:
            summarizer = LexRankSummarizer()
            summary_sentences = summarizer(parser.document, sentence_count)
            summary = " ".join([str(s) for s in summary_sentences])
            return summary
        except Exception as e:
            # Log and fall through to fallback
            current_app.logger.error(f"Sumy summarization failed: {e}")

    # FINAL FALLBACK: simple splitter + heuristic ranking
    sentences = simple_sentence_split(text)
    if not sentences:
        return ""

    # clamp sentence_count
    sentence_count = min(sentence_count, len(sentences))

    # Heuristic: choose longest sentences as a proxy for importance
    ranked = sorted(sentences, key=lambda s: len(s), reverse=True)
    picked = ranked[:sentence_count]

    # Preserve original order for nicer reading
    picked_sorted = sorted(picked, key=lambda s: text.find(s))
    return " ".join(picked_sorted)


@summ_bp.route('/', methods=['GET'])
def form():
    return render_template('summarizer.html')


@summ_bp.route('/process', methods=['POST'])
def process():
    # input can be plain text field 'text' OR uploaded file 'file'
    text_input = request.form.get('text', '').strip()
    try:
        sentences = int(request.form.get('sentences', 5))
    except Exception:
        sentences = 5

    # if file uploaded, prefer file
    if 'file' in request.files and request.files['file'].filename:
        f = request.files['file']
        filename = f.filename.lower()
        ext = filename.rsplit('.', 1)[-1] if '.' in filename else ''
        try:
            if ext == 'pdf':
                f.stream.seek(0)
                extracted = extract_text_from_pdf(f.stream)
            elif ext == 'docx':
                f.stream.seek(0)
                extracted = extract_text_from_docx(f.stream)
            elif ext == 'txt':
                f.stream.seek(0)
                # bytes -> decode
                extracted = f.stream.read().decode('utf-8', errors='ignore')
            else:
                return "Format file tidak didukung.", 415
        except Exception as e:
            current_app.logger.error(f"Error saat ekstrak file: {e}")
            return str(e), 500

        text_input = extracted

    if not text_input:
        return "Tidak ada teks untuk diringkas.", 400

    try:
        summary = summarize_text(text_input, sentence_count=sentences)
        return jsonify({
            "summary": summary,
            "sentences_requested": sentences
        })
    except Exception as e:
        current_app.logger.error(f"Error summarizing: {e}")
        return str(e), 500
